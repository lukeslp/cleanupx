#!/usr/bin/env python3
"""
XCITATION - Citation Extraction and Management Tool

This module provides functionality to extract, deduplicate, and manage citations from
academic papers and documents. It supports:
- Multiple citation styles (APA, MLA, Chicago, IEEE)
- DOI resolution from multiple services
- PDF and document text extraction
- Citation deduplication and validation
- Markdown and JSON output formats

Example:
    >>> from xcitation import process_file_for_citations
    >>> results = process_file_for_citations("paper.pdf", ".citations")
    >>> print(f"Found {results['citations_found']} citations")

Dependencies:
    - requests: For API calls
    - PyPDF2: For PDF text extraction (optional)
    - docx2txt: For DOCX text extraction (optional)
    - python-docx: Alternative DOCX extraction (optional)
"""

import logging
import os
import re
import subprocess
from pathlib import Path
from typing import (
    List, Dict, Optional, Union, Any, Set, TypedDict, Literal,
    NamedTuple, cast
)
import json
from datetime import datetime
from functools import wraps
import time
import requests
from urllib.parse import quote_plus, unquote_plus
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging with more detailed format
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s'
)
logger = logging.getLogger(__name__)

# Type definitions
class Citation(TypedDict, total=False):
    """Type definition for citation data."""
    authors: List[str]
    year: str
    title: Optional[str]
    source: Optional[str]
    doi: Optional[str]
    url: Optional[str]
    volume: Optional[str]
    issue: Optional[str]
    pages: Optional[str]
    type: Literal["doi", "text", "regex"]

class ProcessingResult(TypedDict):
    """Type definition for file processing results."""
    processed: bool
    citations_found: int
    file: str
    error: Optional[str]

class ServiceConfig(NamedTuple):
    """Configuration for DOI resolution services."""
    url: str
    name: str
    requires_key: bool = False
    key_env_var: Optional[str] = None

# Constants
DEFAULT_MODEL_TEXT = "grok-3-mini-latest"
API_BASE_URL = "https://api.x.ai/v1"
MAX_RETRIES = 3
RETRY_DELAY = 2  # seconds
CITATIONS_FILE = "citations.json"
API_TIMEOUT = 30
MAX_CHUNK_SIZE = 50000  # Conservative size to stay well under token limit

# File type constants
TEXT_EXTENSIONS = {
    '.txt', '.md', '.markdown', '.rst', '.text', '.log', '.csv', 
    '.tsv', '.json', '.xml', '.yaml', '.yml', '.html', '.htm',
    '.py', '.db', '.sh', '.rtf', '.ics', '.icsv', '.icsx'
}
DOCUMENT_EXTENSIONS = {'.pdf', '.docx', '.doc', '.ppt', '.pptx', '.xlsx', '.xls'}
IMAGE_EXTENSIONS = {
    '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp',
    '.heic', '.heif', '.ico'
}
MEDIA_EXTENSIONS = {
    '.mp3', '.wav', '.ogg', '.flac', '.aac', '.m4a', '.mp4',
    '.avi', '.mov', '.mkv', '.webm', '.flv', '.m4v'
}
ARCHIVE_EXTENSIONS = {'.zip', '.tar', '.tgz', '.tar.gz', '.rar', '.gz', '.pkg'}

# Supported file extensions for citation extraction
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS.union(DOCUMENT_EXTENSIONS)

# DOI resolution service configurations
DOI_SERVICES = [
    ServiceConfig("https://api.crossref.org/works/", "CrossRef"),
    ServiceConfig("https://doi.org/api/handles/", "DOI.org"),
    ServiceConfig(
        "https://api.unpaywall.org/v2/",
        "Unpaywall",
        requires_key=True,
        key_env_var="UNPAYWALL_API_KEY"
    ),
    ServiceConfig("https://api.datacite.org/works/", "DataCite")
]

# Optional imports with fallbacks
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    logger.debug("pandas not available - Excel file processing will be limited")

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.debug("PyPDF2 not available - PDF processing will use pdftotext only")

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logger.debug("docx2txt not available - trying python-docx")

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    if not DOCX2TXT_AVAILABLE:
        logger.warning("No DOCX processing libraries available")

class CitationError(Exception):
    """Base exception for citation-related errors."""
    pass

class DOIResolutionError(CitationError):
    """Raised when unable to resolve a DOI."""
    pass

class TextExtractionError(CitationError):
    """Raised when unable to extract text from a document."""
    pass

def retry_with_backoff(max_retries: int = MAX_RETRIES, initial_delay: float = RETRY_DELAY):
    """
    Decorator to retry functions with exponential backoff.
    
    Args:
        max_retries: Maximum number of retry attempts
        initial_delay: Initial delay between retries in seconds
        
    Returns:
        Decorated function
    """
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            retries = 0
            current_delay = initial_delay
            
            while retries < max_retries:
                try:
                    return func(*args, **kwargs)
                except (requests.HTTPError, requests.ConnectionError) as e:
                    retries += 1
                    if retries >= max_retries:
                        logger.error(f"Max retries ({max_retries}) reached. Last error: {e}")
                        raise
                    
                    logger.warning(
                        f"API call failed (attempt {retries}/{max_retries}): {e}. "
                        f"Retrying in {current_delay}s..."
                    )
                    time.sleep(current_delay)
                    current_delay *= 2  # Exponential backoff
            
            return func(*args, **kwargs)
        return wrapper
    return decorator

class XAIClient:
    """
    Client for interacting with the X.AI API.
    
    This class provides methods for authenticating and making requests to
    the X.AI API, with built-in error handling and retry logic.
    
    Attributes:
        api_key: X.AI API key
        base_url: Base URL for API requests
        session: Requests session for connection pooling
    """
    
    def __init__(self, api_key: Optional[str] = None):
        """
        Initialize the client with API credentials.
        
        Args:
            api_key: X.AI API key (defaults to XAI_API_KEY environment variable)
            
        Raises:
            ValueError: If no API key is provided or found in environment
        """
        self.api_key = api_key or os.getenv("XAI_API_KEY")
        if not self.api_key:
            raise ValueError("X.AI API key not found. Set XAI_API_KEY environment variable.")
        
        self.base_url = API_BASE_URL
        self.session = requests.Session()
        self.session.headers.update({
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json",
            "User-Agent": "CitationExtractor/1.0 (mailto:support@cleanupx.org)"
        })
    
    @retry_with_backoff()
    def chat(
        self,
        messages: List[Dict[str, str]],
        model: str = DEFAULT_MODEL_TEXT,
        temperature: float = 0.1,
        functions: Optional[List[Dict[str, Any]]] = None
    ) -> Dict[str, Any]:
        """
        Send a chat completion request to the X.AI API.
        
        Args:
            messages: List of message objects with role and content
            model: X.AI model identifier to use
            temperature: Sampling temperature (0-1)
            functions: Optional function calling configuration
            
        Returns:
            Parsed JSON response from the API
            
        Raises:
            requests.HTTPError: If the API request fails
            requests.ConnectionError: If connection fails
            ValueError: If response parsing fails
        """
        url = f"{self.base_url}/chat/completions"
        payload = {
            "model": model,
            "messages": messages,
            "temperature": temperature
        }
        
        # Add function calling if provided
        if functions:
            payload["tools"] = [{"type": "function", "function": f} for f in functions]
            if len(functions) == 1:
                payload["tool_choice"] = {
                    "type": "function", 
                    "function": {"name": functions[0]["name"]}
                }
        
        try:
            logger.debug(f"Sending request to {url}")
            response = self.session.post(url, json=payload, timeout=API_TIMEOUT)
            response.raise_for_status()
            return response.json()
        except requests.HTTPError as e:
            error_details = "No response details available"
            if e.response is not None:
                try:
                    error_details = e.response.json()
                except json.JSONDecodeError:
                    error_details = e.response.text
            
            logger.error(f"API error: {e}\nResponse details: {error_details}")
            raise
    
    def extract_tool_result(self, response: Dict[str, Any]) -> Union[List[Any], Dict[str, Any]]:
        """
        Extract function call result from API response.
        
        Args:
            response: JSON response from chat completion API
            
        Returns:
            Extracted function arguments or content
            
        Raises:
            ValueError: If response parsing fails
        """
        try:
            message = response["choices"][0]["message"]
            
            # Check for tool calls in the response
            if "tool_calls" in message and message["tool_calls"]:
                tool_call = message["tool_calls"][0]
                if tool_call["type"] == "function":
                    return json.loads(tool_call["function"]["arguments"])
            
            # Fallback to content parsing if no tool calls
            content = message.get("content", "")
            # Try to extract JSON from markdown code blocks
            json_match = re.search(r'```(?:json)?\s*([\s\S]*?)\s*```', content, re.DOTALL)
            if json_match:
                try:
                    return json.loads(json_match.group(1))
                except json.JSONDecodeError:
                    pass
            
            # Return the raw content as fallback
            return {"content": content}
            
        except (KeyError, IndexError, json.JSONDecodeError) as e:
            logger.error(f"Error extracting tool result: {e}")
            raise ValueError(f"Failed to parse API response: {e}")

def get_xai_client() -> Optional[XAIClient]:
    """
    Get an initialized X.AI client instance.
    
    Returns:
        XAIClient instance or None if initialization fails
    """
    try:
        return XAIClient()
    except ValueError as e:
        logger.error(f"Failed to initialize XAI client: {e}")
        return None

def extract_citations_from_text(text: str) -> List[Dict[str, Any]]:
    """
    Extract citations from text using the X.AI API.
    
    Args:
        text: Text to analyze
        
    Returns:
        List of citation dictionaries
    """
    client = get_xai_client()
    if not client:
        return []

    # Split text into manageable chunks
    chunks = []
    
    # Split by paragraphs first
    paragraphs = text.split('\n\n')
    current_chunk = ""
    current_size = 0
    
    for para in paragraphs:
        para_size = len(para)
        if current_size + para_size > MAX_CHUNK_SIZE:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = para
            current_size = para_size
        else:
            if current_chunk:
                current_chunk += "\n\n"
            current_chunk += para
            current_size += para_size + 2
    
    if current_chunk:
        chunks.append(current_chunk)
    
    if not chunks:
        chunks = [text]  # Fallback if no paragraphs

    logger.info(f"Split text into {len(chunks)} chunks for processing")
    
    all_citations = []
    seen_citations = set()  # Track duplicates across chunks
    
    for i, chunk in enumerate(chunks, 1):
        logger.info(f"Processing chunk {i}/{len(chunks)} ({len(chunk)} chars)")
        
        messages = [
            {
                "role": "system",
                "content": """Extract academic citations and references.
                Format each citation as a JSON object with:
                - authors (list)
                - year
                - title
                - source
                - doi (if available)
                - url (if available)"""
            },
            {
                "role": "user",
                "content": chunk
            }
        ]

        try:
            response = client.chat(messages, temperature=0.1)  # Lower temperature for more consistent output
            result = client.extract_tool_result(response)
            
            chunk_citations = []
            if isinstance(result, list):
                chunk_citations = result
            elif "content" in result:
                try:
                    parsed = json.loads(result["content"])
                    if isinstance(parsed, list):
                        chunk_citations = parsed
                except json.JSONDecodeError:
                    logger.warning(f"Failed to parse JSON from chunk {i}")
            
            # Deduplicate citations within and across chunks
            for citation in chunk_citations:
                # Create unique key for citation
                if citation.get("doi"):
                    key = clean_doi(citation["doi"])
                else:
                    # For non-DOI citations, use a combination of fields
                    key = f"{citation.get('title', '')}_{citation.get('year', '')}_{','.join(citation.get('authors', []))}"
                
                if key not in seen_citations:
                    all_citations.append(citation)
                    seen_citations.add(key)
            
        except Exception as e:
            logger.error(f"Error processing chunk {i}: {e}")
            if "token" in str(e).lower():
                # Try splitting chunk further if token error
                try:
                    subchunks = [chunk[i:i + MAX_CHUNK_SIZE//2] for i in range(0, len(chunk), MAX_CHUNK_SIZE//2)]
                    logger.info(f"Retrying chunk {i} split into {len(subchunks)} subchunks")
                    for subchunk in subchunks:
                        try:
                            messages[1]["content"] = subchunk
                            response = client.chat(messages, temperature=0.1)
                            result = client.extract_tool_result(response)
                            
                            if isinstance(result, list):
                                for citation in result:
                                    if citation.get("doi"):
                                        key = clean_doi(citation["doi"])
                                    else:
                                        key = f"{citation.get('title', '')}_{citation.get('year', '')}_{','.join(citation.get('authors', []))}"
                                    
                                    if key not in seen_citations:
                                        all_citations.append(citation)
                                        seen_citations.add(key)
                        except Exception as sub_e:
                            logger.error(f"Error processing subchunk of chunk {i}: {sub_e}")
                except Exception as split_e:
                    logger.error(f"Error splitting chunk {i}: {split_e}")
            continue

    # Extract DOIs from text and get their citations
    doi_pattern = r'\b(10\.\d{4,}[/.][^\s/]+)'
    dois = re.findall(doi_pattern, text)
    
    # Get citation info for each unique DOI
    processed_dois = set()
    for doi in dois:
        clean_doi_str = clean_doi(doi)
        if clean_doi_str not in processed_dois:
            citation = extract_citations_from_doi(clean_doi_str)
            if citation:
                key = citation["doi"]
                if key not in seen_citations:
                    all_citations.append(citation)
                    seen_citations.add(key)
                processed_dois.add(clean_doi_str)

    # If no citations found via API or DOIs, try regex
    if not all_citations:
        logger.info("No citations found via API or DOIs, falling back to regex extraction")
        all_citations = _extract_citations_regex(text)

    return all_citations

def _extract_citations_regex(text: str) -> List[Dict[str, Any]]:
    """Extract citations using regex patterns."""
    citations = []
    
    # Common citation patterns
    patterns = [
        # Author (Year) pattern
        r'([A-Z][a-z]+(?:,?\s+(?:&\s+)?[A-Z][a-z]+)*)\s+\((\d{4})\)',
        # (Author, Year) pattern
        r'\(([A-Z][a-z]+(?:,?\s+(?:&\s+)?[A-Z][a-z]+)*),\s+(\d{4})\)',
        # Author et al. (Year) pattern
        r'([A-Z][a-z]+)\s+et\s+al\.\s+\((\d{4})\)',
        # APA style pattern
        r'([A-Za-z\-]+(?:,\s+[A-Z]\.)+(?:\s*&\s*[A-Za-z\-]+(?:,\s+[A-Z]\.)+)*)\s*\((\d{4}[a-z]?)\)',
        # IEEE style pattern
        r'\[(\d+)\]\s+([^.]+)\.\s+"([^"]+)"\s+(?:in\s+)?([^,]+),\s+(?:vol\.\s+)?(\d+)(?:,\s+(?:no\.\s+)?(\d+))?,\s+pp\.\s+(\d+(?:-\d+)?),\s+(\d{4})'
    ]
    
    for pattern in patterns:
        matches = re.finditer(pattern, text)
        for match in matches:
            if '[' in pattern:  # IEEE style
                citation = {
                    "authors": [match.group(2)],
                    "year": match.group(8),
                    "title": match.group(3),
                    "source": match.group(4),
                    "volume": match.group(5),
                    "issue": match.group(6),
                    "pages": match.group(7)
                }
            else:  # Other styles
                citation = {
                    "authors": [author.strip() for author in match.group(1).split('&')],
                    "year": match.group(2),
                    "title": None,
                    "source": None
                }
            citations.append(citation)
    
    return citations

def extract_text_from_pdf(file_path: Union[str, Path], max_pages: int = 3) -> str:
    """
    Extract text from a PDF file using various methods.
    
    Args:
        file_path: Path to the PDF file
        max_pages: Maximum number of pages to process
        
    Returns:
        Extracted text content
    """
    file_path = Path(file_path)
    text = ""
    
    # Try pdftotext first (more reliable)
    try:
        page_option = []
        if max_pages:
            page_option = ["-l", str(max_pages)]
        result = subprocess.run(
            ["pdftotext"] + page_option + ["-layout", str(file_path), "-"],
            capture_output=True,
            text=True,
            check=False
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
        else:
            logger.warning(f"pdftotext extraction failed for {file_path}: {result.stderr}")
    except Exception as e:
        logger.warning(f"pdftotext extraction failed for {file_path}: {e}")
    
    # Fallback to PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f, strict=False)  # Added strict=False to handle malformed PDFs
                pages_to_read = min(len(reader.pages), max_pages) if max_pages else len(reader.pages)
                for i in range(pages_to_read):
                    try:
                        text += reader.pages[i].extract_text() + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {i+1}: {str(e)}")
                        continue
            if text.strip():
                return text
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed for {file_path}: {e}")
    else:
        logger.debug("PyPDF2 not available, skipping PyPDF2 extraction")
    
    return text

def extract_text_from_docx(file_path: Union[str, Path]) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content
    """
    # Try docx2txt first if available
    if DOCX2TXT_AVAILABLE:
        try:
            text = docx2txt.process(file_path)
            if text:
                return text
        except Exception as e:
            logger.warning(f"docx2txt extraction failed for {file_path}: {e}")
    
    # Try python-docx as fallback
    if PYTHON_DOCX_AVAILABLE:
        try:
            doc = Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
            if text:
                return text
        except Exception as e:
            logger.warning(f"python-docx extraction failed for {file_path}: {e}")
    
    logger.error("No available DOCX extraction methods")
    return ""

def extract_text_from_txt(file_path: Union[str, Path]) -> str:
    """
    Extract text from a text file, trying multiple encodings.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        Extracted text content
    """
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                text = f.read()
                if text:
                    return text
        except UnicodeDecodeError:
            continue
        except Exception as e:
            logger.error(f"Error reading {file_path} with {encoding}: {e}")
    
    # If all encodings fail, try binary read as last resort
    try:
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='replace')
    except Exception as e:
        logger.error(f"Binary fallback failed for {file_path}: {e}")
    
    return ""

def extract_text_content(file_path: Union[str, Path]) -> str:
    """
    Extract text content from a file based on its type.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text content
    """
    file_path = Path(file_path)
    ext = file_path.suffix.lower()
    
    try:
        # Use the appropriate extraction method based on file type
        if ext in TEXT_EXTENSIONS:
            return extract_text_from_txt(file_path)
        elif ext in {'.pdf'}:
            return extract_text_from_pdf(file_path)
        elif ext in {'.docx', '.doc'}:
            return extract_text_from_docx(file_path)
        elif ext in {'.xlsx', '.xls'}:
            # For Excel files, try to extract text using pandas
            if PANDAS_AVAILABLE:
                try:
                    df = pd.read_excel(file_path)
                    return df.to_string()
                except Exception as e:
                    logger.error(f"Error processing Excel file {file_path}: {e}")
                    return ""
            else:
                logger.warning("pandas not installed, cannot process Excel files")
                return ""
        else:
            # Fallback to general text extraction
            return extract_text_from_txt(file_path)
            
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {e}")
        return ""

def ensure_metadata_dir(directory: Path) -> Path:
    """
    Ensure the metadata directory exists.
    
    Args:
        directory: Base directory path
        
    Returns:
        Path to the metadata directory
    """
    metadata_dir = directory / ".cleanupx"
    metadata_dir.mkdir(parents=True, exist_ok=True)
    return metadata_dir

def update_citations_file(directory: Path, new_citations: List[Dict[str, Any]]) -> Path:
    """
    Update the citations file in the given directory.
    
    Args:
        directory: Directory path where the citations file should be updated
        new_citations: List of new citations to add
        
    Returns:
        Path to the updated citations file
    """
    citations_file = CITATIONS_FILE(directory)
    citations = []
    
    # Load existing citations if the file exists
    if citations_file.exists():
        try:
            with open(citations_file, 'r', encoding='utf-8') as f:
                citations = json.load(f)
        except Exception as e:
            logger.error(f"Error reading citations file {citations_file}: {e}")
            citations = []
    
    # Add new citations, avoiding duplicates
    existing_raws = {c.get("raw", "") for c in citations}
    for citation in new_citations:
        if citation.get("raw", "") not in existing_raws:
            citation["date_added"] = datetime.now().isoformat()
            citations.append(citation)
            existing_raws.add(citation.get("raw", ""))
    
    # Sort citations by author then year
    citations.sort(key=lambda x: (x.get("authors", ""), x.get("year", "")))
    
    # Write updated citations to file
    try:
        # Ensure .cleanupx directory exists
        citations_file.parent.mkdir(parents=True, exist_ok=True)
        
        with open(citations_file, 'w', encoding='utf-8') as f:
            json.dump(citations, f, indent=2)
        logger.info(f"Updated citations file {citations_file} with {len(new_citations)} new citations")
    except Exception as e:
        logger.error(f"Error writing to citations file {citations_file}: {e}")
    
    return citations_file

def generate_apa_citation_list(directory: Path) -> str:
    """
    Generate a formatted APA citation list from the citations file.
    
    Args:
        directory: Directory path where the citations file is located
        
    Returns:
        Formatted string with APA citations
    """
    citations_file = directory / CITATIONS_FILE
    if not citations_file.exists():
        return "No citations found."
    
    try:
        with open(citations_file, 'r', encoding='utf-8') as f:
            citations = json.load(f)
        
        # Filter to include only reference and DOI citations (not in-text)
        references = [c for c in citations if c.get("type") in ["reference", "doi"]]
        
        # Generate formatted APA citations
        formatted_citations = []
        
        for ref in references:
            authors = ref.get("authors", "")
            year = ref.get("year", "")
            title = ref.get("title", "")
            source = ref.get("source", "")
            volume = ref.get("volume", "")
            issue = ref.get("issue", "")
            pages = ref.get("pages", "")
            doi = ref.get("doi", "")
            url = ref.get("url", "")
            
            # Format the citation in APA style
            citation = f"{authors} ({year}). {title}. "
            
            # Add source information
            if source:
                citation += f"{source}"
                if volume:
                    citation += f", {volume}"
                    if issue:
                        citation += f"({issue})"
                if pages:
                    citation += f", {pages}"
                citation += "."
            
            # Add DOI or URL
            if doi:
                citation += f" https://doi.org/{doi}"
            elif url:
                citation += f" Retrieved from {url}"
            
            formatted_citations.append(citation)
        
        # Join citations with line breaks
        return "\n\n".join(formatted_citations)
        
    except Exception as e:
        logger.error(f"Error generating APA citation list for {directory}: {e}")
        return f"Error generating citations: {e}"

def generate_markdown_citation_list(directory: Path) -> str:
    """
    Generate a markdown-formatted citation list for the directory.
    
    Args:
        directory: Directory path where the citations file is located
        
    Returns:
        Markdown string with APA citations
    """
    citations_file = directory / CITATIONS_FILE
    if not citations_file.exists():
        return "# Citations\n\nNo citations found in this directory."
    
    try:
        with open(citations_file, 'r', encoding='utf-8') as f:
            citations = json.load(f)
        
        # Filter to include only reference and DOI citations (not in-text)
        references = [c for c in citations if c.get("type") in ["reference", "doi"]]
        
        # Generate markdown
        markdown = f"# Citations from {directory.name}\n\n"
        markdown += f"*Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n"
        
        if not references:
            markdown += "No formal citations found in this directory.\n"
            return markdown
        
        markdown += "## References\n\n"
        
        for ref in references:
            authors = ref.get("authors", "")
            year = ref.get("year", "")
            title = ref.get("title", "")
            source = ref.get("source", "")
            volume = ref.get("volume", "")
            issue = ref.get("issue", "")
            pages = ref.get("pages", "")
            doi = ref.get("doi", "")
            url = ref.get("url", "")
            
            # Format the citation in APA style with markdown
            citation = f"{authors} ({year}). *{title}*. "
            
            # Add source information
            if source:
                citation += f"**{source}**"
                if volume:
                    citation += f", {volume}"
                    if issue:
                        citation += f"({issue})"
                if pages:
                    citation += f", {pages}"
                citation += "."
            
            # Add DOI or URL with markdown link
            if doi:
                doi_url = f"https://doi.org/{doi}"
                citation += f" [{doi}]({doi_url})"
            elif url:
                citation += f" [Link]({url})"
            
            markdown += f"- {citation}\n\n"
        
        return markdown
        
    except Exception as e:
        logger.error(f"Error generating markdown citation list for {directory}: {e}")
        return f"# Citations\n\nError generating citations: {e}"

def save_markdown_citations(directory: Path) -> Optional[Path]:
    """
    Save the citations as a markdown file in the directory.
    
    Args:
        directory: Directory path where to save the markdown
        
    Returns:
        Path to the saved markdown file, or None if failed
    """
    try:
        markdown = generate_markdown_citation_list(directory)
        metadata_dir = ensure_metadata_dir(directory)
        output_file = metadata_dir / "CITATIONS.md"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(markdown)
            
        logger.info(f"Saved markdown citations to {output_file}")
        return output_file
    except Exception as e:
        logger.error(f"Error saving markdown citations for {directory}: {e}")
        return None

def process_file_for_citations(
    file_path: Union[str, Path],
    metadata_dir: Path,
    mode: str = "all"
) -> Dict[str, Any]:
    """
    Process a single file for citations.
    
    Args:
        file_path: Path to the file to process
        metadata_dir: Directory to store metadata
        mode: Processing mode ('all', 'new', or 'update')
        
    Returns:
        Dictionary containing processing results
    """
    try:
        file_path = Path(file_path)
        if not file_path.exists():
            return {"error": f"File not found: {file_path}"}
            
        # Check if file type is supported
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            return {"error": f"Unsupported file type: {file_path.suffix}"}
            
        # Get file modification time
        mod_time = datetime.fromtimestamp(file_path.stat().st_mtime)
        
        # Check if we should process this file based on mode
        citations_file = metadata_dir / "citations.json"
        if mode != "all" and citations_file.exists():
            try:
                with open(citations_file, "r", encoding="utf-8") as f:
                    existing = json.load(f)
                    file_key = str(file_path)
                    
                    if file_key in existing:
                        last_processed = datetime.fromisoformat(existing[file_key]["last_processed"])
                        if mode == "new" or (mode == "update" and mod_time <= last_processed):
                            return {"skipped": "Already processed"}
            except Exception as e:
                logger.warning(f"Error reading citations file: {str(e)}")
        
        # Extract content using appropriate method based on file type
        content = ""
        if file_path.suffix.lower() in {'.pdf'}:
            content = extract_text_from_pdf(file_path)
        elif file_path.suffix.lower() in {'.docx', '.doc'}:
            if DOCX2TXT_AVAILABLE:
                content = docx2txt.process(file_path)
            elif PYTHON_DOCX_AVAILABLE:
                doc = Document(file_path)
                content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
            except Exception as e:
                logger.error(f"Error reading file {file_path}: {str(e)}")
                return {"error": f"Failed to read file: {str(e)}"}
                
        if not content:
            return {"error": f"Failed to extract content from {file_path}"}
            
        # Extract citations using API
        citations = extract_citations_from_text(content)
        
        # Look for DOIs and track processed ones
        processed_dois = set()
        doi_pattern = r'\b(10\.\d{4,}[/.][^\s/]+)'
        dois = re.findall(doi_pattern, content)
        
        # Get citation info for each unique DOI
        for doi in dois:
            if doi not in processed_dois:
                citation = extract_citations_from_doi(doi)
                if citation:
                    citations.append(citation)
                    processed_dois.add(doi)
                
        # Deduplicate citations
        unique_citations = []
        seen_citations = set()
        
        for citation in citations:
            # Create a unique key for the citation
            if citation.get("doi"):
                key = clean_doi(citation["doi"])
            else:
                # For non-DOI citations, use a combination of fields
                key = f"{citation.get('title', '')}_{citation.get('year', '')}_{','.join(citation.get('authors', []))}"
            
            if key not in seen_citations:
                unique_citations.append(citation)
                seen_citations.add(key)
                
        # Update citations file
        if unique_citations:
            citations_data = {
                str(file_path): {
                    "citations": unique_citations,
                    "last_processed": mod_time.isoformat(),
                    "file_type": file_path.suffix.lower()
                }
            }
            
            try:
                if citations_file.exists():
                    with open(citations_file, "r", encoding="utf-8") as f:
                        existing = json.load(f)
                    existing.update(citations_data)
                    citations_data = existing
                    
                # Ensure parent directory exists
                citations_file.parent.mkdir(parents=True, exist_ok=True)
                
                with open(citations_file, "w", encoding="utf-8") as f:
                    json.dump(citations_data, f, indent=2)
                    
                # Generate markdown version
                markdown_file = metadata_dir / "citations.md"
                with open(markdown_file, "w", encoding="utf-8") as f:
                    f.write("# Citations\n\n")
                    for file_info in citations_data.values():
                        for citation in file_info["citations"]:
                            authors = citation.get("authors", [])
                            if isinstance(authors, str):
                                authors = [authors]
                            author_text = ", ".join(authors) if authors else "Unknown Author"
                            
                            year = citation.get("year", "n.d.")
                            title = citation.get("title", "Untitled")
                            source = citation.get("source", "")
                            doi = citation.get("doi", "")
                            url = citation.get("url", "")
                            
                            f.write(f"## {title}\n\n")
                            f.write(f"**Authors:** {author_text}\n\n")
                            f.write(f"**Year:** {year}\n\n")
                            if source:
                                f.write(f"**Source:** {source}\n\n")
                            if doi:
                                f.write(f"**DOI:** [{doi}](https://doi.org/{doi})\n\n")
                            if url and url != f"https://doi.org/{doi}":
                                f.write(f"**URL:** {url}\n\n")
                            f.write("---\n\n")
                
            except Exception as e:
                logger.error(f"Error saving citations: {str(e)}")
                return {"error": f"Failed to save citations: {str(e)}"}
            
        return {
            "processed": True,
            "citations_found": len(unique_citations),
            "file": str(file_path)
        }
        
    except Exception as e:
        logger.error(f"Error processing file {file_path}: {str(e)}")
        return {"error": str(e)}

def process_directory(
    directory: Union[str, Path],
    mode: str = "all",
    file_types: Optional[List[str]] = None
) -> Dict[str, Any]:
    """
    Process a directory for citations.
    
    Args:
        directory: Directory to process
        mode: Processing mode ('all', 'new', or 'update')
        file_types: List of file extensions to process (defaults to SUPPORTED_FILE_TYPES)
        
    Returns:
        Dictionary containing processing results
    """
    try:
        directory = Path(directory)
        if not directory.exists():
            return {"error": f"Directory not found: {directory}"}
            
        # Initialize metadata directory
        metadata_dir = ensure_metadata_dir(directory)
        if not metadata_dir:
            return {"error": "Failed to create metadata directory"}
            
        # Get list of files to process
        if not file_types:
            file_types = SUPPORTED_EXTENSIONS
            
        files = []
        for ext in file_types:
            files.extend(directory.rglob(f"*{ext}"))
            
        results = {
            "total_files": len(files),
            "processed": 0,
            "skipped": 0,
            "errors": [],
            "citations_found": 0
        }
        
        # Process each file
        for file_path in files:
            result = process_file_for_citations(file_path, metadata_dir, mode)
            
            if "error" in result:
                results["errors"].append({
                    "file": str(file_path),
                    "error": result["error"]
                })
            elif "skipped" in result:
                results["skipped"] += 1
            else:
                results["processed"] += 1
                results["citations_found"] += result.get("citations_found", 0)
                
        # Generate summary markdown
        summary_file = metadata_dir / "processing_summary.md"
        with open(summary_file, "w", encoding="utf-8") as f:
            f.write("# Citation Processing Summary\n\n")
            f.write(f"**Processed on:** {datetime.now().isoformat()}\n\n")
            f.write(f"**Mode:** {mode}\n\n")
            f.write(f"**Total files found:** {results['total_files']}\n")
            f.write(f"**Files processed:** {results['processed']}\n")
            f.write(f"**Files skipped:** {results['skipped']}\n")
            f.write(f"**Total citations found:** {results['citations_found']}\n\n")
            
            if results["errors"]:
                f.write("## Errors\n\n")
                for error in results["errors"]:
                    f.write(f"- {error['file']}: {error['error']}\n")
                    
        return results
        
    except Exception as e:
        logger.error(f"Error processing directory {directory}: {str(e)}")
        return {"error": str(e)}

def interactive_cli():
    """Interactive command-line interface for citation processing."""
    print("\nWelcome to X.AI Citation Processor!")
    print("This script processes academic files to extract and manage citations.")
    print("Ensure your XAI_API_KEY environment variable is set.")
    
    while True:
        print("\nPlease choose an option:")
        print("  1. Process a directory for citations")
        print("  2. Process a single file for citations")
        print("  3. Generate bibliography in specific format")
        print("  4. Export citations")
        print("  5. Clear citation cache")
        print("  6. Exit")
        
        choice = input("Enter your choice (1-6): ").strip()
        
        if choice == "1":
            directory = input("Enter the full path to the directory: ").strip()
            if not os.path.isdir(directory):
                print(f"Error: {directory} is not a valid directory.")
                continue
            
            mode = input("Enter the mode ('all', 'new', or 'update'): ").strip().lower()
            if mode not in ['all', 'new', 'update']:
                print("Invalid mode. Using 'all'.")
                mode = 'all'
            
            file_types_input = input("Enter file types to process (comma-separated, e.g. '.pdf,.txt', or press Enter for all): ").strip()
            file_types = [ft.strip() for ft in file_types_input.split(',')] if file_types_input else None
            
            results = process_directory(directory, mode=mode, file_types=file_types)
            
            if "error" in results:
                print(f"\nProcessing failed: {results['error']}")
            else:
                print(f"\nProcessing complete!")
                print(f"Total files found: {results['total_files']}")
                print(f"Files processed: {results['processed']}")
                print(f"Files skipped: {results['skipped']}")
                print(f"Total citations found: {results['citations_found']}")
                if results["errors"]:
                    print(f"\nEncountered {len(results['errors'])} errors:")
                    for error in results["errors"]:
                        print(f"- {error['file']}: {error['error']}")
        
        elif choice == "2":
            file_path = input("Enter the full path to the file: ").strip()
            if not os.path.isfile(file_path):
                print(f"Error: {file_path} is not a valid file.")
                continue
            
            verbose_input = input("Enable verbose output? (y/n): ").strip().lower()
            verbose = verbose_input == "y"
            
            result = process_file_for_citations(file_path)
            
            if result["success"]:
                print(f"\nSuccessfully processed: {file_path}")
                print(f"Found {len(result['citations'])} citations")
                if verbose:
                    for citation in result["citations"]:
                        print(f"\n- {citation.get('authors', 'Unknown')} ({citation.get('year', 'Unknown')})")
                        print(f"  {citation.get('title', 'No title')}")
            else:
                print(f"\nError processing {file_path}: {result.get('error', 'Unknown error')}")
        
        elif choice == "3":
            directory = input("Enter the directory path: ").strip()
            if not os.path.isdir(directory):
                print(f"Error: {directory} is not a valid directory.")
                continue
            
            format_choice = input("Choose citation format (apa/mla/chicago/ieee): ").strip().lower()
            if format_choice not in ['apa', 'mla', 'chicago', 'ieee']:
                print("Invalid format. Using APA.")
                format_choice = 'apa'
            
            output = input("Enter output file name (default: bibliography.md): ").strip()
            if not output:
                output = "bibliography.md"
            
            bibliography = create_bibliography(directory, format_choice)
            if bibliography:
                try:
                    with open(os.path.join(directory, output), 'w', encoding='utf-8') as f:
                        f.write(bibliography)
                    print(f"\nBibliography saved to {output}")
                except Exception as e:
                    print(f"Error saving bibliography: {e}")
            else:
                print("No citations found to generate bibliography.")
        
        elif choice == "4":
            directory = input("Enter the directory path: ").strip()
            if not os.path.isdir(directory):
                print(f"Error: {directory} is not a valid directory.")
                continue
            
            format_choice = input("Choose export format (txt/md/bib/csv): ").strip().lower()
            if format_choice not in ['txt', 'md', 'bib', 'csv']:
                print("Invalid format. Using markdown.")
                format_choice = 'md'
            
            output = input("Enter output file name: ").strip()
            if not output:
                output = f"citations.{format_choice}"
            
            citations = generate_citation_list(directory, format_choice)
            if citations:
                try:
                    with open(os.path.join(directory, output), 'w', encoding='utf-8') as f:
                        f.write(citations)
                    print(f"\nCitations exported to {output}")
                except Exception as e:
                    print(f"Error exporting citations: {e}")
            else:
                print("No citations found to export.")
        
        elif choice == "5":
            directory = input("Enter the directory path (or leave blank for current directory): ").strip()
            if not directory:
                directory = os.getcwd()
            
            confirm = input("Are you sure you want to clear the citation cache? (y/n): ").strip().lower()
            if confirm == "y":
                cache_file = os.path.join(directory, CITATIONS_FILE)
                if os.path.exists(cache_file):
                    os.remove(cache_file)
                    print("Citation cache cleared.")
                else:
                    print("No citation cache found.")
        
        elif choice == "6":
            print("Exiting X.AI Citation Processor. Goodbye!")
            break
        
        else:
            print("Invalid choice. Please try again.")

def main():
    """Main entry point for the script."""
    parser = argparse.ArgumentParser(
        description="X.AI Citation Processor: Extract and manage academic citations from files.\n" +
                   "If no arguments are provided, an interactive CLI is launched."
    )
    group = parser.add_mutually_exclusive_group()
    group.add_argument("--directory", help="Path to a directory to process for citations")
    group.add_argument("--file", help="Path to a single file to process")
    
    parser.add_argument("--mode", choices=["all", "new", "update"], default="all",
                      help="Processing mode (all, new, update)")
    parser.add_argument("--format", choices=["apa", "mla", "chicago", "ieee", "txt", "md", "bib", "csv"],
                      help="Output format for citations")
    parser.add_argument("--output", help="Output file path")
    parser.add_argument("--verbose", "-v", action="store_true", help="Enable verbose output")
    parser.add_argument("--recursive", "-r", action="store_true", help="Process subdirectories recursively")
    parser.add_argument("--clear-cache", action="store_true", help="Clear citation cache")
    
    args = parser.parse_args()
    
    if not any([args.directory, args.file, args.clear_cache]):
        interactive_cli()
    else:
        if args.clear_cache:
            directory = args.directory if args.directory else os.getcwd()
            cache_file = os.path.join(directory, CITATIONS_FILE)
            if os.path.exists(cache_file):
                os.remove(cache_file)
                print("Citation cache cleared.")
            else:
                print("No citation cache found.")
            return
        
        if args.file:
            result = process_file_for_citations(args.file)
            if result["success"]:
                print(f"Found {len(result['citations'])} citations")
                if args.verbose:
                    for citation in result["citations"]:
                        print(f"\n- {citation.get('authors', 'Unknown')} ({citation.get('year', 'Unknown')})")
                        print(f"  {citation.get('title', 'No title')}")
            else:
                print(f"Error: {result.get('error', 'Unknown error')}")
        
        if args.directory:
            output_file = args.output if args.output else "citations.md"
            results = process_directory(
                args.directory,
                args.mode,
                args.verbose,
                output_file,
                args.recursive
            )
            
            if results["success"]:
                print(f"Processed {len(results['processed_files'])} files")
                print(f"Found {results['total_citations']} citations")
                if results["errors"]:
                    print(f"Encountered {len(results['errors'])} errors:")
                    for error in results["errors"]:
                        print(f"- {error['file']}: {error['error']}")
            else:
                print(f"Error: {results.get('error', 'Unknown error')}")

if __name__ == "__main__":
    main()